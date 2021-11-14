from docx2python import docx2python
import re
from docx import Document
import os
import pdfplumber
import pandas as pd


def init_docx(fileName):
    file_path = 'Resumes/{}.docx'.format(fileName)
    document = docx2python(file_path)
    doc2 = docx2python(file_path,html=True)
    doc = Document(file_path)
    text = document.text
    lines = [sentences.strip() for sentences in text.split("\n") if len(sentences.strip())>0]

    return lines,doc,document,doc2


def init_pdf(fileName):
    with pdfplumber.open(r"./Resumes/{}.pdf".format(fileName)) as pdf:
        text = pdf.pages[0].extract_text()
        lines = text.split("\n")
        lines = [line.strip() for line in lines]
    return lines,doc


# Returns the educational qualifications from the starting index of educational section
def return_education_points(lines, headingsDict, bold_text_priority,bold_text):
    
    desired_index=None
    next_index = None
    
     
    for headings,index in headingsDict.items():
        next_index = index
        
        # Break the loop when we have both the educational section and the next column's index
        if desired_index is not None:
            break
        
        if re.match("education*",headings,re.I) or re.match("academic*",headings,re.I) or re.search("qualific*",headings,re.I) is not None:
            desired_index = index
            continue
    # Checking for any bold text which is also in capital letters
    if desired_index is None:
        for headings,index in bold_text_priority.items():
            next_index = index
            
            # Break the loop when we have both the educational section and the next column's index
            if desired_index is not None:
                break
            
            if re.match("education*",headings,re.I) or re.match("academic*",headings,re.I) or re.search("qualific*",headings,re.I) is not None:
                desired_index = index
                continue
    
    # Checking for bold text which is not in capitals
    if desired_index is None:
        for headings,index in bold_text.items():
            next_index = index
            
            # Break the loop when we have both the educational section and the next column's index
            if desired_index is not None:
                break
            
            if re.match("education*",headings,re.I) or re.match("academic*",headings,re.I) or re.search("qualific*",headings,re.I) is not None:
                desired_index = index
                continue
        
    # If an educational section exists
    if desired_index is not None:
        
        # If in case the educational section is the last section, print till end of doc, otherwise till next heading
        if next_index!=desired_index:
            return lines[desired_index+1:next_index]
        else:
            return lines[desired_index+1:len(lines)]
        
    else:
        return check_each_line(lines,document)

# Function to determine the headings in the resume
def return_headings(lines):
     headingsDict = {}   
     for lineNo in range(len(lines)):
         try:
             if (re.match("--*",lines[lineNo])==None and re.match("--*",lines[lineNo+1])) or\
             (re.match('\uf0b7+',lines[lineNo]) is None and re.match('\uf0b7+',lines[lineNo+1]) is not None):
                 headingsDict[lines[lineNo].strip()] = lineNo
         except IndexError:
             break
     return headingsDict

# Function to return education lines
def return_lines(lines):
    starting_index = None
    ending_index = None
    for lineNo in range(len(lines)):
        if (check_for_keywords(lines[lineNo]) and check_for_bullets(lines[lineNo+1])):
                starting_index = lineNo+1
                
        if starting_index is not None:
            try:
                temp = lineNo+1
                while(check_for_bullets(lines[temp])):
                    print(check_for_bullets(lines[temp+1]))
                    temp+=1
                ending_index = temp
                        
            except IndexError:
                ending_index = len(lines)
                
            break
            
    
    education_lines = lines[starting_index:ending_index]
    education_lines = list(filter(('\n').__ne__,education_lines))
    education_lines = [line.replace('\n','') for line in education_lines]
    
    return education_lines
                    

# Check for bold text
def return_bold_text(doc,lines):
    
    # Dictionary for bold text
    bold_text = {}
    
    # Dictionary for bold and capital text
    bold_text_priority = {}
    for paragraph in doc.paragraphs:
        for run in paragraph.runs:
            
            # Headings in resumes do not take more than 7 words, doesn't contain special chars and numbers
            if (run.bold or re.sub("[,:]","",run.text).isupper()) and len(run.text.split(" "))<7 and re.search("[0-9]",paragraph.text) is None:
                # Lines with all words in capital have more chances to be in headings
                try:
                    if run.text.isupper():
                        bold_text_priority[run.text]=lines.index(run.text)
                    else:
                        bold_text[run.text]=lines.index(run.text)
                except ValueError:
                    pass      
    # print("Bold Text:-",bold_text)
    # print("Bold Text Priority:-",bold_text_priority)            
    return bold_text_priority,bold_text


# Go through the document paragraph-by-paragraph, take the first line as the heading and check if it marks educational qualifications
def check_with_paragraphs(doc,lines):
    headingsList = []
    continueLoop = True    
    # lines = [sentences.strip() for sentences in doc.text.split("\n") if len(sentences.strip())>0]
    
    for paragraph in doc.paragraphs:
        
        # If we have met the conditions, we will break the loop using continueLoop. 
        if continueLoop:
            for run in paragraph.runs:
                
                if len(headingsList)==2:
                    continueLoop = False
                    break
                
                if len(headingsList)==1 and len(run.text.strip())>0:
                    headingsList.append(run.text)
                    
                if check_for_keywords(run.text):
                    headingsList.append(run.text)
                    
                break
        else:
            break
    
    try:
        starting_index = [i for i, item in enumerate(lines) if re.search("{}*".format(headingsList[0]), item)][0]+1
        ending_index = [i for i, item in enumerate(lines) if re.search("{}*".format(headingsList[1]), item)][0]+1
        
        
    except IndexError:
        if len(headingsList)>0:
            starting_index = [i for i, item in enumerate(lines) if re.search("{}*".format(headingsList[0]), item)][0]+1
            ending_index = len(lines)
        
        else:
            return []
    
    # print("headingsList is:-",headingsList)
    return retLines(lines,starting_index,ending_index)


def retLines(lines,starting_index,ending_index):
    education_lines = lines[starting_index:ending_index-1]
    education_lines = list(filter(('\n').__ne__,education_lines))
    education_lines = [line.replace('\n','') for line in education_lines]
    return education_lines

# Check each line for educational qualifications
def check_each_line(lines,document):
    lines = document.text.splitlines(True)
    lines2 = [line for line in lines if len(line.split(" "))<7]
    education_lines = [line for line in lines2 if check_for_keywords(line)]
    
    try:
        starting_index = lines.index(education_lines[0])+1
        ending_index = len(lines)
        
        for x in range(starting_index,ending_index):
            try:
                if lines[x+1]=='\n' and lines[x+2]=='\n':
                    ending_index = x+2
                    
            except IndexError:
                pass

        return retLines(lines,starting_index,ending_index-1)

    except IndexError:
        return ["404"]


def return_lines_using_wordCount(lines):
    
    lineNumber = [lineNo for lineNo in range(len(lines)) if check_for_keywords(lines[lineNo])][0]
    temp = lineNumber+1
    
    try:
        
        # Check if there are more than 3 words in a line or number marking cgpa/percentage/gpa etc.
        while(len(lines[temp].split(" "))>=3 or re.fullmatch("[0-9]+\.*[0-9]* *%* *[A-z]*",lines[temp]) is not None):
            temp+=1
    except IndexError:
        pass
    return retLines(lines,lineNumber+1,temp+1)

def check_for_keywords(headings):
    if re.match("education*",headings,re.I) or re.match("ac?dem*",headings,re.I) or re.search("qualific*",headings,re.I) is not None:
        return True
    else: False
    
    
def check_for_bullets(text):
    if re.match("--*",text) is not None or re.match('\uf0b7+',text) is not None:
                    return True
                        
    else: return False

# Formatting and refining the output
def format_points(education_points,*charReplacements):
    
    for replacement in charReplacements:
        education_points = [points.replace(replacement,"") for points in education_points]
    return education_points
    

#%% This cell is just used to print stuff
# print("\n\n",format_points(return_education_points(return_headings(lines),return_bold_text())))

if __name__ == '__main__':
    
    # Using dataframes to carry the content of all resumes
    df = pd.DataFrame(columns=['Name','Qualifications'])
    
    cwd = os.getcwd()
    getCurrentFileNames = os.listdir(cwd+"/Resume/")
    
    for files in getCurrentFileNames:
        fileName = files.split(".")
        try:
            
            # If the document is a docx document
            if fileName[1]=='docx':
                lines,doc,document,doc2 = init_docx(fileName[0])
                bold_text_priority,bold_text = return_bold_text(doc,lines)
                content_to_be_written = return_education_points(lines,return_headings(lines),bold_text_priority,bold_text)
                
                # If the extracted information contains more than just the educational information required.
                if len(content_to_be_written)<=20:
                    content_to_be_written = format_points(content_to_be_written,"--\\t","\t")
                else:
                    # print("in else with",fileName[0])
                    
                    content_to_be_written = format_points(check_with_paragraphs(doc,lines))
                    
                    if len(content_to_be_written)>20:
                        content_to_be_written = format_points(return_lines(lines))
                    else:
                        content_to_be_written = format_points(content_to_be_written,"--\\t","\t")
                    
            # If the document is a pdf document
            else:
                lines,doc = init_pdf(fileName[0])
                content_to_be_written = format_points(return_education_points(lines,return_headings(lines),bold_text_priority,bold_text),"\uf0b7")
                
            df.loc[len(df.index)] = [fileName[0],content_to_be_written ]
                
            print(return_lines_using_wordCount(lines))
            
            
            # The code to write the output in a text file
            with open(cwd+"/Output/"+fileName[0]+".txt","w") as f:
                for content in content_to_be_written:
                    f.write(content+"\n")
        
        except FileNotFoundError:
            pass
            
            
#%% Checking for individual resumes (only for testing purposes)

# Finding problem resumes

problemResumes = []
df.apply(lambda x: problemResumes.append(x['Name']) if (len(x['Qualifications'])>20) else None,axis=1)

print("Problem Resumes Ratio:-",len(problemResumes)*100/len(df.index))


# print(df.loc[df['Name']=='Anuj Kumar']['Qualifications'])
            
#%% Failure count

series_of_failures = df['Qualifications'].apply(lambda x: 1 if '404' in x or x==[] else 0).tolist()
print("Ares of failures",series_of_failures.count(0)*100/len(series_of_failures))



#%% Test Cell











